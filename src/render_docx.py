"""Render tailored CV HTML to an ATS-safe .docx.

Manual-only, and an optional extra: python-docx is deliberately kept out of
requirements.txt so the cron pipeline's stdlib+pyyaml dependency posture is
unchanged. Install it only if you need DOCX:

    pip install python-docx

Why DOCX at all: some application portals parse .docx reliably and mangle or
reject PDFs, particularly PDFs produced by headless-browser printing, where
text can come out as positioned runs rather than a clean reading order.

"ATS-safe" here means the boring subset that every parser handles: built-in
heading styles, plain paragraphs, real bullet lists, no tables, no text
boxes, no headers/footers, no columns. Anything cleverer risks a parser
reading the document out of order or dropping content entirely.
"""

from html.parser import HTMLParser

# Tags whose text becomes a paragraph of the mapped style. Anything not listed
# is treated as body text rather than dropped — losing a line silently would
# be worse than styling it plainly.
_STYLE_FOR_TAG = {
    "h1": "Title",
    "h2": "Heading 1",
    "h3": "Heading 2",
    "h4": "Heading 3",
    "li": "List Bullet",
}
_BLOCK_TAGS = set(_STYLE_FOR_TAG) | {"p", "div"}


class _Block:
    def __init__(self, style):
        self.style = style
        self.runs = []  # (text, bold)

    def text(self):
        return "".join(t for t, _ in self.runs).strip()


class _CVHTMLParser(HTMLParser):
    """Flatten body HTML into an ordered list of styled blocks.

    Only bold is carried through as inline formatting — it's the one emphasis
    that survives ATS parsing intact and the only one the tailor prompt emits.
    """

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocks = []
        self._current = None
        self._bold_depth = 0
        self._skip_depth = 0

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style"):
            self._skip_depth += 1
        elif tag in ("strong", "b"):
            self._bold_depth += 1
        elif tag in _BLOCK_TAGS:
            self._flush()
            self._current = _Block(_STYLE_FOR_TAG.get(tag))
        elif tag == "br" and self._current is not None:
            self._current.runs.append(("\n", False))

    def handle_endtag(self, tag):
        if tag in ("script", "style") and self._skip_depth:
            self._skip_depth -= 1
        elif tag in ("strong", "b") and self._bold_depth:
            self._bold_depth -= 1
        elif tag in _BLOCK_TAGS:
            self._flush()

    def handle_data(self, data):
        if self._skip_depth or not data.strip():
            return
        if self._current is None:
            # Text outside any block tag still belongs in the document.
            self._current = _Block(None)
        self._current.runs.append((data, self._bold_depth > 0))

    def _flush(self):
        if self._current is not None and self._current.text():
            self.blocks.append(self._current)
        self._current = None

    def close(self):
        super().close()
        self._flush()


def parse_blocks(body_html):
    parser = _CVHTMLParser()
    parser.feed(body_html or "")
    parser.close()
    return parser.blocks


def render_docx(body_html, out_path):
    """Write body_html to out_path as .docx. Raises SystemExit if python-docx
    isn't installed, since this is a manual path where a clear instruction
    beats a traceback."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise SystemExit(
            "DOCX output needs python-docx, which is an optional extra "
            "(kept out of requirements.txt so the cron stays stdlib-only).\n"
            "Install it with:  pip install python-docx"
        )

    document = Document()
    # 11pt Calibri is the Word default and the safest possible choice — an
    # embedded or unusual font is a common cause of garbled ATS extraction.
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for block in parse_blocks(body_html):
        paragraph = document.add_paragraph(style=block.style) if block.style else document.add_paragraph()
        for text, bold in block.runs:
            if not text.strip():
                continue
            run = paragraph.add_run(text)
            run.bold = bold

    document.save(out_path)
    return out_path
